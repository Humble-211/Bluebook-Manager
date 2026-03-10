"""
Bluebook Manager — Sharing service (share quality documents across bluebooks).
"""

import os
import re

from config import SHAREABLE_SECTIONS, STORAGE_ROOT
from dal import dal
from dal.models import Bluebook, BluebookFile
from services.log_service import log


def _append_die_to_docx(file_path: str, target_die_number: str):
    """Append a target die number to 'Die:' patterns inside a DOCX file.

    For example, 'Die: 12345' becomes 'Die: 12345-67890'.
    Works across paragraphs and table cells.
    """
    abs_path = os.path.join(STORAGE_ROOT, file_path)
    if not abs_path.lower().endswith(".docx") or not os.path.isfile(abs_path):
        return

    try:
        from docx import Document
        doc = Document(abs_path)
        # Pattern: "Die:" followed by optional space and die numbers (with optional dashes)
        pattern = re.compile(r'(Die:\s*[\w\-]+)')

        def update_text(text):
            """Append target die number to Die: patterns if not already present."""
            def replacer(match):
                current = match.group(1)
                # Check if target die is already appended
                if target_die_number in current:
                    return current
                return f"{current}-{target_die_number}"
            return pattern.sub(replacer, text)

        modified = False

        # Process paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                new_text = update_text(run.text)
                if new_text != run.text:
                    run.text = new_text
                    modified = True

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            new_text = update_text(run.text)
                            if new_text != run.text:
                                run.text = new_text
                                modified = True

        if modified:
            doc.save(abs_path)
            log("SHARE_FILE_DOCX_UPDATE",
                f"Appended Die# {target_die_number} to {file_path}")

    except Exception as e:
        log("SHARE_FILE_DOCX_ERROR",
            f"Failed to update DOCX {file_path}: {e}")


def share_file(file_id: int, target_bluebook_ids: list[int]):
    """Share a file to one or more target bluebooks."""
    bf = dal.get_bluebook_file(file_id)
    if not bf:
        raise ValueError(f"File id={file_id} not found.")

    if bf.section_type not in SHAREABLE_SECTIONS:
        raise ValueError(f"Section '{bf.section_type}' does not support sharing. "
                         f"Only {SHAREABLE_SECTIONS} can be shared.")

    source_bb = dal.get_bluebook(bf.bluebook_id)
    shared_count = 0

    for target_id in target_bluebook_ids:
        # Don't share to the file's own bluebook
        if target_id == bf.bluebook_id:
            continue
        target_bb = dal.get_bluebook(target_id)
        if not target_bb:
            continue
        dal.add_shared_file(file_id, target_id)
        shared_count += 1

        # Update DOCX content with target die number
        _append_die_to_docx(bf.file_path, target_bb.die_number)

        log("SHARE_FILE",
            f"File '{bf.file_path}' from Die# {source_bb.die_number} "
            f"shared to Die# {target_bb.die_number}")

    return shared_count


def unshare_file_from_bluebook(file_id: int, bluebook_id: int):
    """Remove a shared file reference from a single bluebook."""
    dal.remove_shared_file(file_id, bluebook_id)
    log("UNSHARE_FILE", f"file_id={file_id} removed from bluebook_id={bluebook_id}")


def unshare_file_from_all(file_id: int):
    """Remove a shared file from all target bluebooks."""
    dal.remove_all_shared_refs(file_id)
    log("UNSHARE_FILE_ALL", f"file_id={file_id} removed from all targets")


def get_shared_targets(file_id: int) -> list[Bluebook]:
    """Get list of bluebooks this file is shared to."""
    return dal.get_shared_targets(file_id)


def is_file_shared(file_id: int) -> bool:
    return dal.is_file_shared(file_id)
