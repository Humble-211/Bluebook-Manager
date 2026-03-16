"""
Bluebook Manager — File service (create, attach, open, remove files).
"""

import os
import shutil

import pythoncom
import win32com.client

from config import (
    SECTION_FILE_TYPES,
    SECTION_FOLDERS,
    STORAGE_ROOT,
    TEMPLATE_DIR,
    TEMPLATE_SECTIONS,
)
from dal import dal
from dal.models import BluebookFile
from services.log_service import log


def resolve_shortcut(path: str) -> str:
    """If *path* is a .lnk shortcut, return its target; otherwise return *path* unchanged."""
    if not path.lower().endswith(".lnk"):
        return path
    try:
        pythoncom.CoInitialize()
        shell = win32com.client.Dispatch("WScript.Shell")
        shortcut = shell.CreateShortCut(path)
        target = shortcut.TargetPath
        return target if target else path
    except Exception:
        return path


def create_shortcut(shortcut_path: str, target_path: str):
    """Create a Windows .lnk shortcut at *shortcut_path* pointing to *target_path*."""
    pythoncom.CoInitialize()
    shell = win32com.client.Dispatch("WScript.Shell")
    shortcut = shell.CreateShortCut(shortcut_path)
    shortcut.TargetPath = target_path
    shortcut.WorkingDirectory = os.path.dirname(target_path)
    shortcut.save()


def get_section_folder(die_number: str, section_type: str) -> str:
    """Get the absolute disk path for a bluebook section folder."""
    return os.path.join(STORAGE_ROOT, die_number, SECTION_FOLDERS[section_type])


def create_from_template(bluebook_id: int, section_type: str,
                         filename: str) -> BluebookFile:
    """Create a new file from a template and register it."""
    if section_type not in TEMPLATE_SECTIONS:
        raise ValueError(f"Section '{section_type}' does not support template creation.")

    bb = dal.get_bluebook(bluebook_id)
    if not bb:
        raise ValueError(f"Bluebook id={bluebook_id} not found.")

    # Determine template file
    template_map = {
        "cover": "cover_template.docx",
        "quality_alerts": "quality_alert_template.docx",
        "quality_notes": "quality_notes_template.docx",
        "packing_instruction": "packing_instruction_template.docx",
        "fit_and_functions": "fit_and_functions_template.docx",
    }
    template_path = os.path.join(TEMPLATE_DIR, template_map[section_type])
    if not os.path.isfile(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    # Ensure filename has .docx extension
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    # Copy template to section folder
    dest_folder = get_section_folder(bb.die_number, section_type)
    os.makedirs(dest_folder, exist_ok=True)
    dest_path = os.path.join(dest_folder, filename)

    if os.path.exists(dest_path):
        raise FileExistsError(f"File already exists: {dest_path}")

    shutil.copy2(template_path, dest_path)

    # Auto-fill table fields for quality alerts and quality notes
    if section_type in ("quality_alerts", "quality_notes"):
        _autofill_template(dest_path, bb, filename)

    # Store relative path in DB
    rel_path = os.path.relpath(dest_path, STORAGE_ROOT)
    order = dal.get_next_display_order(bluebook_id, section_type)
    fid = dal.add_bluebook_file(bluebook_id, section_type, rel_path, order)
    log("CREATE_FILE", f"Bluebook Die# {bb.die_number}, section={section_type}, file={filename}")

    return dal.get_bluebook_file(fid)


def _autofill_template(dest_path: str, bb, filename: str = ""):
    """Auto-fill the first table in a Quality Alert/Notes DOCX with bluebook info.

    Scans table cells for labels like 'Customer', 'Die', 'Type of Complaint',
    'Date', and 'Q.A #' and fills in the corresponding values.
    """
    import re
    from datetime import datetime

    try:
        from docx import Document
        doc = Document(dest_path)

        if not doc.tables:
            return

        customer_names = ", ".join(bb.customer_names) if bb.customer_names else ""
        die_number = bb.die_number
        today = datetime.now().strftime("%m/%d/%Y")
        complaint_type = f"{customer_names} REJECTED" if customer_names else "REJECTED"

        # Extract QA number from filename (e.g. QA-26-014 from QA-26-014-15901-Screw-Hole.docx)
        qa_number = ""
        m = re.match(r"^QA-\d{2}-0?(\d{1,3})", filename, re.IGNORECASE)
        if m:
            qa_number = m.group(1)

        table = doc.tables[0]
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                text = cell.text.strip().lower()

                # Find label cells and fill the adjacent cell (or same cell after colon)
                if 'customer' in text and ':' in cell.text:
                    _fill_cell_value(cell, "Customer", customer_names)
                elif 'die' in text and ':' in cell.text:
                    _fill_cell_value(cell, "Die", die_number)
                elif 'type of complaint' in text and ':' in cell.text:
                    _fill_cell_value(cell, "Type of Complaint", complaint_type)
                elif 'date' in text and ':' in cell.text:
                    _fill_cell_value(cell, "Date", today)

        table1 = doc.tables[1]
        for row in table1.rows:
            for i, cell in enumerate(row.cells):
                text = cell.text.strip().lower()
                if 'q.a' in text and ':' in cell.text:
                    if qa_number:
                        _fill_cell_value(cell, "Q.A", qa_number)
        doc.save(dest_path)
    except Exception as e:
        log("AUTOFILL_ERROR", f"Could not auto-fill template {dest_path}: {e}")


def _fill_cell_value(cell, label: str, value: str):
    """Replace the content of a cell that has 'Label:' format with 'Label: value'."""
    for para in cell.paragraphs:
        for run in para.runs:
            # Find the label pattern and replace
            if label.lower() in run.text.lower() and ':' in run.text:
                # Keep everything up to and including the colon, then add value
                colon_idx = run.text.index(':')
                prefix = run.text[:colon_idx + 1]
                run.text = f"{prefix} {value}"
                return
    # Fallback: if no run matched, try setting the first paragraph
    if cell.paragraphs:
        for para in cell.paragraphs:
            if label.lower() in para.text.lower() and ':' in para.text:
                colon_idx = para.text.index(':')
                prefix = para.text[:colon_idx + 1]
                # Capture formatting from the first run before clearing
                is_bold = None
                font_name = None
                font_size = None
                if para.runs:
                    first_run = para.runs[0]
                    is_bold = first_run.bold
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                para.clear()
                new_run = para.add_run(f"{prefix} {value}")
                if is_bold is not None:
                    new_run.bold = is_bold
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                return


def rename_file(file_id: int, new_filename: str):
    """Rename a file on disk and update its path in the database."""
    bf = dal.get_bluebook_file(file_id)
    if not bf:
        raise ValueError(f"File id={file_id} not found.")

    old_abs = get_absolute_path(bf.file_path)
    if not os.path.isfile(old_abs):
        raise FileNotFoundError(f"File not found on disk: {old_abs}")

    # For shortcuts, preserve the .lnk extension
    is_lnk = old_abs.lower().endswith(".lnk")
    if is_lnk and not new_filename.lower().endswith(".lnk"):
        new_filename = new_filename + ".lnk"

    new_abs = os.path.join(os.path.dirname(old_abs), new_filename)

    if os.path.exists(new_abs):
        raise FileExistsError(f"A file named '{new_filename}' already exists.")

    try:
        os.rename(old_abs, new_abs)
    except PermissionError:
        raise PermissionError(
            f"Cannot rename '{os.path.basename(old_abs)}' because it is "
            f"currently open in another program.\n\n"
            f"Please close the file and try again.")

    new_rel = os.path.relpath(new_abs, STORAGE_ROOT)
    dal.update_bluebook_file_path(file_id, new_rel)
    log("RENAME_FILE", f"file_id={file_id}, old={bf.file_path}, new={new_rel}")


def attach_shortcut(bluebook_id: int, section_type: str,
                    source_path: str) -> BluebookFile:
    """Create a .lnk shortcut in local storage pointing to *source_path* and register it."""
    bb = dal.get_bluebook(bluebook_id)
    if not bb:
        raise ValueError(f"Bluebook id={bluebook_id} not found.")

    if not os.path.isfile(source_path):
        raise FileNotFoundError(f"Source file not found: {source_path}")

    dest_folder = get_section_folder(bb.die_number, section_type)
    os.makedirs(dest_folder, exist_ok=True)

    filename = os.path.basename(source_path)
    lnk_name = filename + ".lnk"
    lnk_path = os.path.join(dest_folder, lnk_name)

    # Handle name collision
    base = filename  # e.g. "15901.pdf"
    counter = 1
    while os.path.exists(lnk_path):
        lnk_name = f"{os.path.splitext(base)[0]}_{counter}{os.path.splitext(base)[1]}.lnk"
        lnk_path = os.path.join(dest_folder, lnk_name)
        counter += 1

    create_shortcut(lnk_path, source_path)

    rel_path = os.path.relpath(lnk_path, STORAGE_ROOT)
    order = dal.get_next_display_order(bluebook_id, section_type)
    fid = dal.add_bluebook_file(bluebook_id, section_type, rel_path, order)
    log("ATTACH_SHORTCUT", f"Bluebook Die# {bb.die_number}, section={section_type}, "
                           f"shortcut={lnk_name} -> {source_path}")

    return dal.get_bluebook_file(fid)


def attach_file(bluebook_id: int, section_type: str,
                source_path: str) -> BluebookFile:
    """Attach an external file to a bluebook section.

    For master_drawings, creates a shortcut (.lnk) instead of copying.
    For all other sections, copies the file into local storage.
    """
    # Sections that use shortcuts instead of copying
    if section_type in ("master_drawings", "qc_drawings"):
        # Validate file type before delegating
        ext = os.path.splitext(source_path)[1].lower()
        allowed = SECTION_FILE_TYPES.get(section_type, [])
        if ext not in allowed:
            raise ValueError(f"File type '{ext}' not allowed for section '{section_type}'. "
                             f"Allowed: {allowed}")
        return attach_shortcut(bluebook_id, section_type, source_path)

    bb = dal.get_bluebook(bluebook_id)
    if not bb:
        raise ValueError(f"Bluebook id={bluebook_id} not found.")

    # Validate file type
    ext = os.path.splitext(source_path)[1].lower()
    allowed = SECTION_FILE_TYPES.get(section_type, [])
    if ext not in allowed:
        raise ValueError(f"File type '{ext}' not allowed for section '{section_type}'. "
                         f"Allowed: {allowed}")

    if not os.path.isfile(source_path):
        raise FileNotFoundError(f"Source file not found: {source_path}")

    dest_folder = get_section_folder(bb.die_number, section_type)
    os.makedirs(dest_folder, exist_ok=True)
    filename = os.path.basename(source_path)
    dest_path = os.path.join(dest_folder, filename)

    # Handle name collision
    base, extension = os.path.splitext(filename)
    counter = 1
    while os.path.exists(dest_path):
        dest_path = os.path.join(dest_folder, f"{base}_{counter}{extension}")
        counter += 1

    shutil.copy2(source_path, dest_path)

    rel_path = os.path.relpath(dest_path, STORAGE_ROOT)
    order = dal.get_next_display_order(bluebook_id, section_type)
    fid = dal.add_bluebook_file(bluebook_id, section_type, rel_path, order)
    log("ATTACH_FILE", f"Bluebook Die# {bb.die_number}, section={section_type}, "
                       f"file={os.path.basename(dest_path)}")

    return dal.get_bluebook_file(fid)


def remove_file(file_id: int, delete_from_disk: bool = True):
    """Remove a file record and optionally delete from disk."""
    bf = dal.get_bluebook_file(file_id)
    if not bf:
        return

    if delete_from_disk:
        abs_path = get_absolute_path(bf.file_path)
        # Only delete files that live in our local storage
        if os.path.isfile(abs_path) and abs_path.startswith(STORAGE_ROOT):
            try:
                os.remove(abs_path)
            except PermissionError:
                fname = os.path.basename(abs_path)
                raise PermissionError(
                    f"Cannot delete '{fname}' because it is currently open "
                    f"in another program.\n\n"
                    f"Please close the file and try again.")

    dal.delete_bluebook_file(file_id)
    log("DELETE_FILE", f"file_id={file_id}, path={bf.file_path}, disk={delete_from_disk}")


def open_file(file_path: str):
    """Open a file in the default Windows application.

    Handles both relative and absolute paths, and resolves .lnk shortcuts.
    """
    abs_path = get_absolute_path(file_path)
    if not os.path.isfile(abs_path):
        raise FileNotFoundError(f"File not found: {abs_path}")

    # Resolve shortcut to its target
    target = resolve_shortcut(abs_path)
    if target != abs_path and not os.path.isfile(target):
        raise FileNotFoundError(
            f"Shortcut target not found: {target}\n"
            f"The file on the network drive may have been moved or deleted.")

    os.startfile(target)
    log("OPEN_FILE", f"path={file_path}")


def get_absolute_path(file_path: str) -> str:
    """Convert a relative file_path (from DB) to absolute."""
    if os.path.isabs(file_path):
        return file_path
    return os.path.join(STORAGE_ROOT, file_path)


def get_files_for_section(bluebook_id: int, section_type: str) -> list[BluebookFile]:
    """Get all files (own + shared) for a specific section."""
    return dal.get_files_for_bluebook(bluebook_id, section_type=section_type)


def get_all_files(bluebook_id: int) -> list[BluebookFile]:
    """Get all files for a bluebook."""
    return dal.get_files_for_bluebook(bluebook_id)


