import os
import re
import sys

# Ensure project root is importable and config.BASE_DIR resolves correctly.
# config.py uses sys.argv[0] to set BASE_DIR, so we override it before import.
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)
sys.argv[0] = os.path.join(PROJECT_ROOT, "main.py")

from docx import Document

from dal.database import init_db
from dal import dal
from services import file_service

# ── Configuration ──
SCAN_DIR = r"S:\QCLAB\Bluebooks Related\Fit-Functions\Word"
SECTION_TYPE = "fit_and_functions"

# Pattern to capture die number(s) after "Die:" / "Die#" / "Die :"
# Matches: "Die: 15901", "Die# 15899/15901", "Die: 15925b", etc.
# Stops before unrelated text like "Colour:"
DIE_PATTERN = re.compile(
    r"Die\s*[:#]?\s*(\d+[A-Za-z]?(?:\s*[/,\-]\s*\d+[A-Za-z]?)*)",
    re.IGNORECASE,
)


def extract_die_raw(filepath: str) -> str | None:
    """Open a .docx file and return the raw text after 'Die:' ."""
    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"  ⚠  Could not open: {filepath} ({e})")
        return None

    # Search paragraphs
    for para in doc.paragraphs:
        m = DIE_PATTERN.search(para.text)
        if m:
            return m.group(1).strip()

    # Also search inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                m = DIE_PATTERN.search(cell.text)
                if m:
                    return m.group(1).strip()

    return None


def parse_die_numbers(raw: str) -> list[str]:
    """Split a raw die string into individual die numbers.

    Handles separators: /  ,  -
    Strips trailing letters: 15925b → 15925
    """
    # Split on / , or -
    parts = re.split(r"[/,\-]", raw)

    dies = []
    for part in parts:
        part = part.strip()
        if not part or not any(ch.isdigit() for ch in part):
            continue  # skip words with no digits
        # Strip trailing letters (e.g. 15925b → 15925)
        cleaned = re.sub(r"[A-Za-z]+$", "", part).strip()
        if cleaned:
            dies.append(cleaned)

    return dies


def main():
    if not os.path.isdir(SCAN_DIR):
        print(f"ERROR: Scan directory not found: {SCAN_DIR}")
        sys.exit(1)

    init_db()

    # Collect all .docx files
    docx_files = []
    for root, _dirs, files in os.walk(SCAN_DIR):
        for f in files:
            if f.lower().endswith(".docx") and not f.startswith("~$"):
                docx_files.append(os.path.join(root, f))

    if not docx_files:
        print(f"No .docx files found in {SCAN_DIR}")
        return

    print(f"Found {len(docx_files)} .docx file(s) to process.\n")

    linked = []
    failed = []

    for filepath in docx_files:
        filename = os.path.basename(filepath)
        raw_die = extract_die_raw(filepath)

        if not raw_die:
            failed.append((filepath, "No 'Die:' field found in document"))
            continue

        die_numbers = parse_die_numbers(raw_die)
        if not die_numbers:
            failed.append((filepath, f"Could not parse die number(s) from: '{raw_die}'"))
            continue

        file_linked = False
        for die_number in die_numbers:
            # Look up the bluebook by exact die number
            bb = dal.get_bluebook_by_die(die_number)
            if not bb:
                failed.append((filepath, f"No bluebook found for Die# {die_number}"))
                continue

            # Check if file is already attached (by matching filename)
            existing_files = dal.get_files_for_bluebook(bb.id, SECTION_TYPE)
            already_attached = any(
                os.path.basename(ef.file_path) == filename for ef in existing_files
            )
            if already_attached:
                print(f"  ⊘  {filename}  →  Die# {die_number} (already attached)")
                file_linked = True
                continue

            # Attach the file
            try:
                bf = file_service.attach_file(bb.id, SECTION_TYPE, filepath)
                linked.append((filepath, die_number, bf.id))
                file_linked = True
                print(f"  ✓  {filename}  →  Die# {die_number}")
            except Exception as e:
                failed.append((filepath, f"Error attaching to Die# {die_number}: {e}"))

    # ── Summary ──
    print(f"\n{'='*60}")
    print(f"  Linked:  {len(linked)}")
    print(f"  Failed:  {len(failed)}")
    print(f"{'='*60}")

    if failed:
        print("\n⚠  Files that could NOT be linked:\n")
        for path, reason in failed:
            print(f"  ✗  {path}")
            print(f"     Reason: {reason}\n")


if __name__ == "__main__":
    main()
